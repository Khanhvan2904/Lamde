import streamlit as st
from docx import Document
import io
import random
import re

# --- HÀM HỖ TRỢ XỬ LÝ DOCX ---
def get_questions_from_docx(file):
    """Chia nhỏ nội dung file Word thành danh sách các câu hỏi dựa trên chữ 'Câu'"""
    doc = Document(file)
    questions = []
    current_q = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        # Nếu gặp chữ "Câu [số]" thì bắt đầu câu mới
        if re.match(r'^Câu\s*\d+', text, re.IGNORECASE):
            if current_q:
                questions.append(current_q)
            current_q = [para] # Lưu cả đối tượng paragraph để giữ định dạng
        else:
            if current_q:
                current_q.append(para)
                
    if current_q:
        questions.append(current_q)
    return questions

def create_docx_output(selected_questions, version_code):
    """Tạo file Word mới từ danh sách câu hỏi đã chọn"""
    new_doc = Document()
    new_doc.add_heading(f'MÃ ĐỀ THI: {version_code}', 0)
    
    global_q_num = 1
    for q_group in selected_questions:
        for i, para in enumerate(q_group):
            new_p = new_doc.add_paragraph()
            # Copy nội dung và định dạng đơn giản
            text = para.text
            if i == 0: # Dòng đầu tiên của câu hỏi
                text = re.sub(r'^Câu\s*\d+', f'Câu {global_q_num}', text, flags=re.IGNORECASE)
            new_p.text = text
        global_q_num += 1
        new_doc.add_paragraph("") # Khoảng trống giữa các câu
        
    buffer = io.BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- GIAO DIỆN STREAMLIT ---
st.set_page_config(page_title="Hệ thống Tạo Đề Ngẫu Nhiên", layout="wide")
st.title("🎯 Hệ thống Trích xuất & Tạo Đề Thi Tự Động")

with st.sidebar:
    st.header("⚙️ Cấu hình đề thi")
    num_versions = st.number_input("Số lượng mã đề cần tạo", min_value=1, max_value=20, value=1)
    
    st.divider()
    n_multi = st.number_input("Số câu Trắc nghiệm nhiều lựa chọn", min_value=0, value=12)
    n_tf = st.number_input("Số câu Trắc nghiệm Đúng/Sai", min_value=0, value=4)
    n_short = st.number_input("Số câu Trắc nghiệm Trả lời ngắn", min_value=0, value=6)

# Giao diện Upload 3 vùng riêng biệt
col1, col2, col3 = st.columns(3)

with col1:
    st.subheader("1. Nhiều lựa chọn")
    file_multi = st.file_uploader("Upload ngân hàng Dạng 1", type=["docx"], key="multi")

with col2:
    st.subheader("2. Đúng/Sai")
    file_tf = st.file_uploader("Upload ngân hàng Dạng 2", type=["docx"], key="tf")

with col3:
    st.subheader("3. Trả lời ngắn")
    file_short = st.file_uploader("Upload ngân hàng Dạng 3", type=["docx"], key="short")

# --- LOGIC XỬ LÝ CHÍNH ---
if st.button("🚀 Bắt đầu tạo đề thi", type="primary"):
    if file_multi and file_tf and file_short:
        # Bước 1: Trích xuất câu hỏi từ 3 nguồn
        bank_multi = get_questions_from_docx(file_multi)
        bank_tf = get_questions_from_docx(file_tf)
        bank_short = get_questions_from_docx(file_short)
        
        # Kiểm tra số lượng
        if len(bank_multi) < n_multi or len(bank_tf) < n_tf or len(bank_short) < n_short:
            st.error("❌ Số lượng câu hỏi trong ngân hàng không đủ so với yêu cầu!")
        else:
            st.success(f"✅ Đã tải: {len(bank_multi)} câu TN, {len(bank_tf)} câu Đ/S, {len(bank_short)} câu ngắn.")
            
            # Bước 2: Tạo từng mã đề
            for v in range(num_versions):
                v_code = 101 + v
                
                # Lấy ngẫu nhiên theo số lượng yêu cầu
                selected = (
                    random.sample(bank_multi, n_multi) +
                    random.sample(bank_tf, n_tf) +
                    random.sample(bank_short, n_short)
                )
                
                # Bước 3: Build file docx
                docx_file = create_docx_output(selected, v_code)
                
                st.download_button(
                    label=f"📥 Tải xuống Mã đề {v_code}",
                    data=docx_file,
                    file_name=f"Ma_De_{v_code}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.warning("⚠️ Vui lòng upload đầy đủ cả 3 tệp ngân hàng câu hỏi.")